"""
Smart CV Optimizer - ATS-style CV Analysis Tool

This Streamlit application helps users optimize their CVs to pass Applicant Tracking Systems (ATS)
by analyzing alignment with Job Descriptions using LLM APIs (GPT-4o or Gemini 1.5 Pro).

Key Features:
- PDF text extraction with caching
- LLM-powered ATS simulation and analysis
- Match score calculation (0-100%)
- Keyword gap analysis
- Bullet point optimization (X-Y-Z formula)
- ATS red flag detection
"""

import streamlit as st
import pdfplumber
import json
import hashlib
import time
import re
from io import BytesIO
from datetime import datetime
from dateutil import parser as date_parser
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# ============================================================================
# CONSTANTS
# ============================================================================

# Token control limits (character counts to prevent excessive API costs)
MAX_CV_LENGTH = 12000  # ~3000 tokens
MAX_JD_LENGTH = 6000   # ~1500 tokens

# API retry configuration
MAX_RETRIES = 3
RETRY_DELAYS = [1, 2, 4]  # Exponential backoff in seconds

# API timeout
API_TIMEOUT = 60  # seconds

# ============================================================================
# PDF TEXT EXTRACTION
# ============================================================================

@st.cache_data(show_spinner=False)
def extract_text_from_pdf(pdf_bytes):
    """
    Extract text from PDF file using pdfplumber.
    
    Uses Streamlit caching to avoid reprocessing the same PDF on reruns.
    The cache key is based on the PDF content (bytes), so different PDFs
    will be processed separately.
    
    Args:
        pdf_bytes: PDF file content as bytes
        
    Returns:
        Extracted text as string, or None if extraction fails
    """
    try:
        # Create a BytesIO object from bytes for pdfplumber
        pdf_file = BytesIO(pdf_bytes)
        
        # Extract text from all pages
        text_content = []
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        # Join all pages with newlines
        full_text = "\n".join(text_content)
        
        # Normalize whitespace: remove excessive newlines and spaces
        # Replace multiple newlines with double newline
        full_text = "\n\n".join([line.strip() for line in full_text.split("\n") if line.strip()])
        
        return full_text
    
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return None


def detect_vertical_whitespace_river(page):
    """
    Dynamically detect the vertical whitespace 'river' that separates columns.
    
    This function analyzes character bounding boxes to find the largest continuous
    vertical gap that could represent a column separator. NO hardcoded percentages.
    
    Logic:
    1. Extract all character bounding boxes on the page
    2. Analyze x-coordinates to find largest continuous vertical gap
    3. Gap must span at least 60% of page height to be valid
    4. Return center of river as dynamic split_x coordinate
    5. If no significant river found, return None (1-column layout)
    
    Args:
        page: pdfplumber Page object
        
    Returns:
        float: x-coordinate of column separator, or None if 1-column
    """
    chars = page.chars
    if not chars:
        return None
    
    # Sort characters by x-coordinate to find gaps
    x_coords = sorted(set(char['x0'] for char in chars))
    
    if len(x_coords) < 10:  # Too few x-coordinates, probably 1-column
        return None
    
    # Find largest gap between consecutive x-coordinates
    max_gap = 0
    split_x = None
    
    for i in range(len(x_coords) - 1):
        gap = x_coords[i + 1] - x_coords[i]
        
        if gap > max_gap and gap > 20:  # Minimum 20 points gap
            # Check if this gap spans 60% of page height (vertical river test)
            left_x = x_coords[i]
            right_x = x_coords[i + 1]
            
            # Get characters on left and right of this gap
            left_chars = [c for c in chars if c['x1'] <= left_x]
            right_chars = [c for c in chars if c['x0'] >= right_x]
            
            if left_chars and right_chars:
                # Calculate vertical span on each side
                left_span = max(c['y1'] for c in left_chars) - min(c['y0'] for c in left_chars)
                right_span = max(c['y1'] for c in right_chars) - min(c['y0'] for c in right_chars)
                
                page_height = page.height
                
                # Both sides must span at least 60% of page height
                if left_span >= 0.6 * page_height and right_span >= 0.6 * page_height:
                    max_gap = gap
                    split_x = (left_x + right_x) / 2
    
    return split_x


@st.cache_data(show_spinner=False)
def extract_text_with_layout(pdf_bytes):
    """
    Extract text with intelligent 2-column layout detection using river detection.
    
    This function dynamically detects column separators instead of using hardcoded
    percentages. Works with narrow, wide, or no sidebars.
    
    Args:
        pdf_bytes: PDF file content as bytes
        
    Returns:
        dict: {
            "sidebar": text or None,
            "main": text,
            "header": text or None,
            "layout": "1-col" or "2-col",
            "split_x": float or None
        }
    """
    try:
        pdf_file = BytesIO(pdf_bytes)
        
        with pdfplumber.open(pdf_file) as pdf:
            all_sidebar = []
            all_main = []
            layout_type = None
            split_x_coord = None
            
            for page in pdf.pages:
                # Detect column split dynamically
                split_x = detect_vertical_whitespace_river(page)
                
                if split_x is None:
                    # 1-column layout - extract normally
                    full_text = page.extract_text()
                    if full_text:
                        all_main.append(full_text)
                    layout_type = "1-col"
                else:
                    # 2-column layout - crop and extract separately
                    layout_type = "2-col"
                    split_x_coord = split_x
                    
                    # Left column (sidebar)
                    left_bbox = (0, 0, split_x, page.height)
                    sidebar_text = page.crop(left_bbox).extract_text()
                    if sidebar_text:
                        all_sidebar.append(sidebar_text)
                    
                    # Right column (main content)
                    right_bbox = (split_x, 0, page.width, page.height)
                    main_text = page.crop(right_bbox).extract_text()
                    if main_text:
                        all_main.append(main_text)
            
            # Join all pages
            sidebar_combined = "\n\n".join(all_sidebar) if all_sidebar else None
            main_combined = "\n\n".join(all_main) if all_main else ""
            
            # Normalize whitespace
            if sidebar_combined:
                sidebar_combined = "\n\n".join([line.strip() for line in sidebar_combined.split("\n") if line.strip()])
            main_combined = "\n\n".join([line.strip() for line in main_combined.split("\n") if line.strip()])
            
            return {
                "sidebar": sidebar_combined,
                "main": main_combined,
                "header": None,
                "layout": layout_type or "1-col",
                "split_x": split_x_coord
            }
    
    except Exception as e:
        st.error(f"Error extracting text with layout: {str(e)}")
        # Fallback to simple extraction
        simple_text = extract_text_from_pdf(pdf_bytes)
        return {
            "sidebar": None,
            "main": simple_text or "",
            "header": None,
            "layout": "1-col",
            "split_x": None
        }


# ============================================================================
# TEXT PROCESSING
# ============================================================================

def truncate_text(text, max_length):
    """
    Truncate text to maximum length while preserving word boundaries.
    
    This prevents excessive token usage and API costs. We truncate at word
    boundaries to avoid cutting words in half, which could confuse the LLM.
    
    Args:
        text: Input text string
        max_length: Maximum character length
        
    Returns:
        Truncated text with ellipsis if truncated
    """
    if len(text) <= max_length:
        return text
    
    # Find the last space before max_length to avoid cutting words
    truncated = text[:max_length]
    last_space = truncated.rfind(" ")
    
    if last_space > 0:
        truncated = truncated[:last_space]
    
    return truncated + "..."


def generate_cache_key(cv_text, jd_text, llm_choice):
    """
    Generate a unique cache key for LLM responses.
    
    Uses SHA-256 hash of the inputs to create a deterministic cache key.
    This allows us to cache LLM responses and avoid duplicate API calls
    when the user hasn't changed any inputs.
    
    Args:
        cv_text: CV text content
        jd_text: Job description text
        llm_choice: Selected LLM model
        
    Returns:
        Hash string to use as cache key
    """
    cache_string = f"{cv_text}|{jd_text}|{llm_choice}"
    return hashlib.sha256(cache_string.encode()).hexdigest()


def fix_future_dates(cv_text):
    """
    Automatically detect and fix future dates in CV text.
    
    Fixes logical errors like "July 2025 - Present" when current year is 2024.
    Preserves "Present" and "Current" keywords.
    
    Args:
        cv_text: Original CV text
        
    Returns:
        tuple: (corrected_text, list_of_corrections)
    """
    corrections = []
    current_year = datetime.now().year
    current_month = datetime.now().month
    
    # Pattern 1: "Month YYYY" format
    def replace_month_year(match):
        month_name = match.group(1)
        year_str = match.group(2)
        year = int(year_str)
        
        if year > current_year:
            # Future year - replace with current year
            corrected = f"{month_name} {current_year}"
            corrections.append(f"'{month_name} {year}' → '{corrected}'")
            return corrected
        
        return match.group(0)
    
    # Pattern 2: Standalone "YYYY" (but not in ranges handled by pattern 3)
    month_names = r'(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)'
    
    # Replace "Month YYYY" patterns
    pattern1 = re.compile(rf'\b({month_names})\s+(\d{{4}})\b', re.IGNORECASE)
    corrected_text = pattern1.sub(replace_month_year, cv_text)
    
    # Pattern 3: "YYYY - Present" or "YYYY - Current"
    def replace_year_range(match):
        year_str = match.group(1)
        year = int(year_str)
        separator = match.group(2)
        end_word = match.group(3)
        
        if year > current_year:
            corrected = f"{current_year}{separator}{end_word}"
            corrections.append(f"'{year}{separator}{end_word}' → '{corrected}'")
            return corrected
        
        return match.group(0)
    
    pattern3 = re.compile(r'\b(\d{4})(\s*[-–—]\s*)(Present|Current|present|current)\b')
    corrected_text = pattern3.sub(replace_year_range, corrected_text)
    
    return corrected_text, corrections


def validate_quantifiable_bullets(revised_cv_text):
    """
    Validate that revised bullets contain quantifiable metrics (X-Y-Z formula).
    
    X-Y-Z Formula requires measurable results: "Accomplished [X] as measured by [Y], by doing [Z]"
    The [Y] component MUST contain a metric: percentage, dollar amount, multiplier, time, etc.
    
    Args:
        revised_cv_text: Revised CV text to validate
        
    Returns:
        list: Tuples of (bullet_text, has_metric, warning_message)
    """
    # Regex patterns for quantifiable metrics
    metric_patterns = [
        r'\b\d+%',                                           # Percentages: 25%, 100%
        r'\$\d+[\d,]*(?:\.\d+)?[KMB]?',                     # Currency: $50K, $1.5M, $100
        r'\b\d+x\b',                                         # Multipliers: 3x, 10x
        r'\b\d+\s*(?:hours?|days?|weeks?|months?|years?)',  # Time: 40 hours, 3 months
        r'\b\d+[\d,]*\s*(?:users?|customers?|clients?|transactions?|requests?|records?|files?|items?)',  # Counts
        r'\b(?:increased|decreased|reduced|improved|grew|boosted|enhanced|optimized)\s+(?:by\s+)?\d+',  # Change metrics
        r'\b\d+[\d,]*\+\b',                                  # Plus notation: 1000+, 50K+
        r'\b\d+[\d,]*\s*(?:million|billion|thousand|[KMB])\b',  # Large numbers: 5 million, 10K
    ]
    
    combined_pattern = '|'.join(metric_patterns)
    
    # Extract bullets from revised CV
    lines = revised_cv_text.split('\n')
    bullets = []
    
    for line in lines:
        stripped = line.strip()
        # Identify bullet points (various markers)
        if stripped and any(stripped.startswith(marker) for marker in ['-', '•', '*', '–', '◦']) and len(stripped) > 20:
            bullet_text = stripped.lstrip('-•*–◦ ').strip()
            
            # Check if bullet contains any metric
            has_metric = bool(re.search(combined_pattern, bullet_text, re.IGNORECASE))
            
            if not has_metric:
                # Flag bullet as needing metric
                warning = f"⚠️ [NEEDS METRIC] {bullet_text}"
                bullets.append((bullet_text, False, warning))
            else:
                bullets.append((bullet_text, True, None))
    
    return bullets


# ============================================================================
# SCORE VALIDATION
# ============================================================================

def validate_and_clamp_scores(data):
    """
    Validate and clamp all scores to their valid ranges.
    
    This is critical for reliability - LLMs sometimes return values outside
    the expected range. We clamp scores to prevent displaying invalid data:
    - match_score: 0-100
    - technical_skills: 0-50
    - experience_alignment: 0-30
    - keyword_coverage: 0-20
    
    Args:
        data: Analysis results dictionary from LLM
        
    Returns:
        Validated data with clamped scores
    """
    # Clamp main match score to 0-100
    if "match_score" in data:
        data["match_score"] = max(0, min(100, data["match_score"]))
    
    # Clamp score breakdown components
    if "score_breakdown" in data:
        breakdown = data["score_breakdown"]
        
        if "technical_skills" in breakdown:
            breakdown["technical_skills"] = max(0, min(50, breakdown["technical_skills"]))
        
        if "experience_alignment" in breakdown:
            breakdown["experience_alignment"] = max(0, min(30, breakdown["experience_alignment"]))
        
        if "keyword_coverage" in breakdown:
            breakdown["keyword_coverage"] = max(0, min(20, breakdown["keyword_coverage"]))
    
    return data


# ============================================================================
# LLM API INTEGRATION
# ============================================================================

def create_revision_prompt():
    """
    Create the system prompt for CV revision.
    
    This prompt instructs the LLM to:
    1. Preserve all original structure (headings, sections, bullet points)
    2. Apply X-Y-Z formula with quantifiable metrics
    3. Fix logical date errors (future dates)
    4. Naturally integrate missing keywords from analysis
    5. Maintain template-specific formatting (Skills groups, Project structure)
    6. Return complete revised CV text
    
    Returns:
        Revision prompt string
    """
    return """You are a world-class career coach, ATS expert, professional resume writer, and CV optimization specialist.

Your task is to revise a CV to improve its ATS compatibility and alignment with a Job Description while STRICTLY preserving its original structure.

CRITICAL RULES - STRUCTURE PRESERVATION:
1. **DO NOT add or remove sections** - Keep all original headings exactly as they are
2. **DO NOT add or remove bullet points** - Maintain the exact number of bullets in each section
3. **DO NOT change the order of sections** - Keep the original sequence
4. **DO NOT change formatting style** - Preserve the layout approach (bullets, paragraphs, etc.)

TEMPLATE-SPECIFIC PRESERVATION (If Present):
- Keep Skills groups separate: "Tools & Platforms", "Frontend", "Backend", "Databases", "Languages"
- Preserve Project structure: Bold Title → Tech Stack line → Bullet points
- Maintain Contact information format: Email, Phone, Location

DATE VALIDATION - CRITICAL:
- **Fix ALL future dates** (e.g., if a date shows "July 2025" and current year is 2024, change it to "July 2024")
- Check for logical errors in employment dates
- "Present" or "Current" are acceptable end dates
- Example fix: "Family Builders | July 2025 - Present" → "Family Builders | July 2024 - Present"

X-Y-Z FORMULA APPLICATION:
- Apply the formula: "Accomplished [X] as measured by [Y], by doing [Z]"
- **[Y] MUST contain quantifiable metrics**: percentages, dollar amounts, time saved, counts
- Use strong action verbs: "Architected", "Spearheaded", "Optimized", "Engineered" (not "Worked on", "Helped")
- Examples:
  - ✅ "Improved API response time by 45% through Redis caching, reducing average latency from 800ms to 440ms"
  - ✅ "Reduced cloud infrastructure costs by $50K annually via containerization and auto-scaling"
  - ❌ "Worked on improving system performance" (too vague, no metrics)

WHAT YOU MUST DO:
1. **Fix date errors**: Correct any future dates or illogical date ranges
2. **Apply X-Y-Z formula**: Rewrite bullets with quantifiable metrics where possible
3. **Improve wording and grammar**: Fix errors, enhance clarity, use stronger action verbs
4. **Integrate missing keywords**: Naturally weave in relevant keywords from the analysis where appropriate
5. **Enhance impact statements**: Make achievements more compelling and measurable
6. **Maintain professional tone**: Keep it concise, recruiter-friendly, and ATS-compatible
7. **Preserve the original voice**: Don't drastically change the writing style

OUTPUT FORMAT:
Return the complete revised CV text with:
- All original section headings preserved
- Same number of bullet points per section
- Fixed dates (no future dates unless explicitly correct)
- Improved wording with X-Y-Z formula and metrics
- Professional formatting ready for use

DO NOT include explanations, comments, or meta-text. Output ONLY the revised CV content."""


def create_system_prompt():
    """
    Create the system prompt for LLM analysis.
    
    This prompt is carefully engineered to:
    1. Simulate ATS parsing behavior (not human recruiter behavior)
    2. Enforce structured JSON output only
    3. Prevent hallucinations (no invented skills/metrics)
    4. Follow exact scoring formula
    5. Apply ATS-specific rules (keyword matching, formatting checks)
    
    Returns:
        System prompt string
    """
    return """You are an expert Applicant Tracking System (ATS) analyzer and career coach. Your task is to simulate how an ATS system parses and scores a CV against a Job Description.

CRITICAL RULES:
1. **ATS Simulation Mode**: Behave like an ATS parser, not a human recruiter
   - Emphasize keyword frequency and exact matches (ATS systems are literal)
   - Penalize missing or non-standard section headers (e.g., "Professional Experience" vs "Experience")
   - Flag formatting issues: multi-column layouts, tables, icons, graphics
   
2. **Structured JSON Output**: Return ONLY valid JSON, no free-form text

3. **Scoring Formula** (must follow exactly):
   - Technical skills match: 0-50 points (50% weight)
   - Years of experience alignment: 0-30 points (30% weight)
   - Keyword coverage & relevance: 0-20 points (20% weight)
   - Total match_score: 0-100 (sum of above, clamped to 0-100)

4. **Keyword Extraction**: Extract keywords ONLY from the Job Description
   - Do NOT invent or infer skills not explicitly mentioned in the JD
   - Do NOT add keywords present in CV but absent from JD
   
5. **No Fabrication**: 
   - Never add fake metrics, tools, or experience
   - If metrics are missing, use placeholders like "measurable impact" or "improved efficiency"
   - Preserve original meaning of bullet points

6. **X-Y-Z Formula**: Rewrite bullets as "Accomplished [X] as measured by [Y], by doing [Z]"

Return ONLY this JSON structure:
{
  "match_score": <0-100 integer>,
  "score_breakdown": {
    "technical_skills": <0-50 integer>,
    "experience_alignment": <0-30 integer>,
    "keyword_coverage": <0-20 integer>
  },
  "score_explanation": "<Brief explanation of how score was calculated>",
  "keyword_gaps": {
    "required_skills": ["<skill from JD missing/weak in CV>"],
    "preferred_skills": ["<preferred skill from JD missing/weak in CV>"],
    "tools_technologies": ["<tool/tech from JD missing/weak in CV>"],
    "soft_skills": ["<soft skill from JD missing/weak in CV>"]
  },
  "bullet_optimizations": [
    {
      "original": "<original bullet from CV>",
      "optimized": "<rewritten using X-Y-Z formula>",
      "reason": "<why this improves ATS matching>"
    }
  ],
  "ats_red_flags": [
    {
      "issue": "<ATS parsing issue detected>",
      "explanation": "<why ATS systems struggle with this>",
      "fix": "<actionable fix>"
    }
  ]
}"""


def call_openai_api(cv_text, jd_text, max_retries=MAX_RETRIES):
    """
    Call OpenAI GPT-4o API for CV analysis.
    
    Uses JSON mode to enforce structured output and includes retry logic
    for handling transient failures or invalid JSON responses.
    
    Args:
        cv_text: CV text content
        jd_text: Job description text
        max_retries: Maximum number of retry attempts
        
    Returns:
        Parsed JSON response or None if all retries fail
    """
    try:
        from openai import OpenAI
        
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        user_prompt = f"""Analyze this CV against the Job Description.

JOB DESCRIPTION:
{jd_text}

CV:
{cv_text}

Provide analysis in JSON format as specified."""
        
        for attempt in range(max_retries):
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": create_system_prompt()},
                        {"role": "user", "content": user_prompt}
                    ],
                    response_format={"type": "json_object"},
                    timeout=API_TIMEOUT
                )
                
                # Parse JSON response
                result = json.loads(response.choices[0].message.content)
                
                # Validate and clamp scores
                result = validate_and_clamp_scores(result)
                
                return result
            
            except json.JSONDecodeError as e:
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAYS[attempt])
                    continue
                else:
                    st.error(f"Failed to parse valid JSON after {max_retries} attempts")
                    return None
            
            except Exception as e:
                error_str = str(e)
                # Check for quota/billing errors (don't retry these)
                if "quota" in error_str.lower() or "insufficient_quota" in error_str.lower() or "429" in error_str:
                    # Check if it's specifically a billing/payment method issue
                    is_billing_issue = "insufficient_quota" in error_str.lower() and "billing" in error_str.lower()
                    
                    if is_billing_issue or "insufficient_quota" in error_str.lower():
                        st.error("""
                        **OpenAI API Quota Error** ❌
                        
                        **Most Common Cause:** No payment method added to your account.
                        
                        Even if your usage shows $0.00, OpenAI requires a payment method on file to use the API.
                        
                        **Quick Fix (2 minutes):**
                        1. **Add Payment Method** (REQUIRED):
                           👉 https://platform.openai.com/account/billing/payment-methods
                           - Add a credit/debit card
                           - Even if you have free credits, payment method is required
                        
                        2. **Set Usage Limits** (Optional but recommended):
                           - Visit: https://platform.openai.com/account/billing/limits
                           - Set a monthly spending limit to control costs
                        
                        3. **Verify Billing Status:**
                           - Check: https://platform.openai.com/account/billing
                           - Ensure account is active
                        
                        **Alternative: Use Gemini (Free Tier Available)**
                        - Switch to "Gemini 1.5 Pro" in the sidebar
                        - Get free API key: https://makersuite.google.com/app/apikey
                        - No payment method required for free tier
                        """)
                    else:
                        st.error("""
                        **OpenAI API Quota Exceeded** ❌
                        
                        You've exceeded your current OpenAI API quota. Here's how to fix it:
                        
                        1. **Check your quota & billing:**
                           - Visit: https://platform.openai.com/usage
                           - Check: https://platform.openai.com/account/billing
                        
                        2. **Solutions:**
                           - Add payment method: https://platform.openai.com/account/billing/payment-methods
                           - Upgrade your plan if needed
                           - Wait for quota reset (usually monthly)
                           - **OR switch to Gemini 1.5 Pro** (often has free tier)
                        
                        3. **Use Gemini instead:**
                           - Select "Gemini 1.5 Pro" in the sidebar
                           - Get your free API key: https://makersuite.google.com/app/apikey
                        """)
                    return None
                
                # For other errors, retry if attempts remain
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAYS[attempt])
                    continue
                else:
                    st.error(f"OpenAI API error: {error_str}")
                    return None
    
    except ImportError:
        st.error("OpenAI library not installed. Run: pip install openai")
        return None
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {str(e)}")
        return None


def call_gemini_api(cv_text, jd_text, max_retries=MAX_RETRIES):
    """
    Call Google Gemini API for CV analysis.
    
    Uses JSON mode and retry logic similar to OpenAI implementation.
    Tries multiple model names for compatibility across different API versions.
    
    Args:
        cv_text: CV text content
        jd_text: Job description text
        max_retries: Maximum number of retry attempts
        
    Returns:
        Parsed JSON response or None if all retries fail
    """
    try:
        import google.generativeai as genai
        
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        
        # Try multiple model names for compatibility
        # Different API versions/regions may have different model names
        model_names = [
            "gemini-1.5-pro",      # Latest version
            "gemini-1.5-flash",    # Faster alternative
            "gemini-pro",           # Standard version
            "models/gemini-1.5-pro",  # With models/ prefix
            "models/gemini-pro"     # With models/ prefix
        ]
        
        model = None
        model_name_used = None
        
        # Try to find an available model
        for name in model_names:
            try:
                test_model = genai.GenerativeModel(
                    model_name=name,
                    generation_config={
                        "temperature": 0.1,
                        "response_mime_type": "application/json"
                    }
                )
                # Test if model is accessible
                model = test_model
                model_name_used = name
                break
            except Exception:
                continue
        
        # If no model worked, try listing available models
        if model is None:
            try:
                available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                if available_models:
                    # Use the first available model that looks like gemini
                    gemini_models = [m for m in available_models if 'gemini' in m.lower()]
                    if gemini_models:
                        model_name_used = gemini_models[0]
                        model = genai.GenerativeModel(
                            model_name=model_name_used,
                            generation_config={
                                "temperature": 0.1,
                                "response_mime_type": "application/json"
                            }
                        )
                    else:
                        # Fallback to first available model
                        model_name_used = available_models[0]
                        model = genai.GenerativeModel(
                            model_name=model_name_used,
                            generation_config={
                                "temperature": 0.1,
                                "response_mime_type": "application/json"
                            }
                        )
                else:
                    st.error("No Gemini models available. Please check your API key and region.")
                    return None
            except Exception as e:
                st.error(f"Could not list available models: {str(e)}")
                st.info("Trying with 'gemini-pro' as fallback...")
                model_name_used = "gemini-pro"
                model = genai.GenerativeModel(
                    model_name=model_name_used,
                    generation_config={
                        "temperature": 0.1,
                        "response_mime_type": "application/json"
                    }
                )
        
        # Final check - ensure we have a valid model
        if model is None:
            st.error("Could not initialize any Gemini model. Please check your API key and try again.")
            return None
        
        user_prompt = f"""{create_system_prompt()}

Analyze this CV against the Job Description.

JOB DESCRIPTION:
{jd_text}

CV:
{cv_text}

Provide analysis in JSON format as specified."""
        
        for attempt in range(max_retries):
            try:
                response = model.generate_content(user_prompt)
                
                # Parse JSON response
                result = json.loads(response.text)
                
                # Validate and clamp scores
                result = validate_and_clamp_scores(result)
                
                return result
            
            except json.JSONDecodeError as e:
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAYS[attempt])
                    continue
                else:
                    st.error(f"Failed to parse valid JSON after {max_retries} attempts")
                    return None
            
            except Exception as e:
                error_str = str(e)
                # Check for model not found errors (don't retry these)
                if "404" in error_str or "not found" in error_str.lower() or "not supported" in error_str.lower():
                    model_display = model_name_used if model_name_used else "selected model"
                    st.error(f"""
                    **Gemini Model Not Found** ❌
                    
                    The model '{model_display}' is not available. This could be due to:
                    
                    1. **API Version/Region Issue:**
                       - Different regions may have different model names
                       - Try updating your API key or checking region settings
                    
                    2. **Model Name Changed:**
                       - Google may have updated model names
                       - Check available models: https://ai.google.dev/models/gemini
                    
                    3. **API Key Permissions:**
                       - Ensure your API key has access to Gemini models
                       - Get a new key: https://makersuite.google.com/app/apikey
                    
                    **Troubleshooting:**
                    - The app tried multiple model names automatically
                    - If this persists, check Google AI Studio for current model names
                    - Error details: {error_str}
                    """)
                    return None
                
                # For other errors, retry if attempts remain
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAYS[attempt])
                    continue
                else:
                    st.error(f"Gemini API error: {error_str}")
                    return None
    
    except ImportError:
        st.error("Google Generative AI library not installed. Run: pip install google-generativeai")
        return None
    except Exception as e:
        st.error(f"Error initializing Gemini client: {str(e)}")
        return None


def call_llm_api(cv_text, jd_text, llm_choice):
    """
    Main LLM API wrapper with caching.
    
    Routes to appropriate API based on user selection and caches results
    in session state to prevent duplicate API calls when inputs haven't changed.
    
    Args:
        cv_text: CV text content
        jd_text: Job description text
        llm_choice: Selected LLM ("GPT-4o" or "Gemini 1.5 Pro")
        
    Returns:
        Analysis results dictionary or None if API call fails
    """
    # Check cache first (prevents duplicate API calls)
    cache_key = generate_cache_key(cv_text, jd_text, llm_choice)
    
    if "llm_cache" not in st.session_state:
        st.session_state.llm_cache = {}
    
    if cache_key in st.session_state.llm_cache:
        return st.session_state.llm_cache[cache_key]
    
    # Truncate inputs for token control (cost optimization)
    cv_text_truncated = truncate_text(cv_text, MAX_CV_LENGTH)
    jd_text_truncated = truncate_text(jd_text, MAX_JD_LENGTH)
    
    # Show warning if truncation occurred
    if len(cv_text) > MAX_CV_LENGTH:
        st.warning(f"⚠️ CV text was truncated from {len(cv_text)} to {MAX_CV_LENGTH} characters to control API costs.")
    if len(jd_text) > MAX_JD_LENGTH:
        st.warning(f"⚠️ Job Description was truncated from {len(jd_text)} to {MAX_JD_LENGTH} characters to control API costs.")
    
    # Call appropriate API
    result = None
    with st.spinner(f"Analyzing with {llm_choice}..."):
        if llm_choice == "GPT-4o":
            result = call_openai_api(cv_text_truncated, jd_text_truncated)
        elif llm_choice == "Gemini 1.5 Pro":
            result = call_gemini_api(cv_text_truncated, jd_text_truncated)
    
    # Cache result if successful
    if result:
        st.session_state.llm_cache[cache_key] = result
    
    return result


def call_revision_llm(cv_text, jd_text, analysis_results, llm_choice):
    """
    Call LLM to generate revised CV with integrated feedback.
    
    Routes to appropriate API based on user selection and caches results
    in session state to prevent duplicate API calls.
    
    Args:
        cv_text: Original CV text content
        jd_text: Job description text
        analysis_results: Dictionary with analysis feedback (keyword gaps, etc.)
        llm_choice: Selected LLM ("GPT-4o" or "Gemini 1.5 Pro")
        
    Returns:
        Revised CV text string or None if API call fails
    """
    # Check cache first
    cache_key = hashlib.sha256(f"revision_{cv_text}_{jd_text}_{llm_choice}".encode()).hexdigest()
    
    if "revision_cache" not in st.session_state:
        st.session_state.revision_cache = {}
    
    if cache_key in st.session_state.revision_cache:
        return st.session_state.revision_cache[cache_key]
    
    # Truncate inputs for token control
    cv_text_truncated = truncate_text(cv_text, MAX_CV_LENGTH)
    jd_text_truncated = truncate_text(jd_text, MAX_JD_LENGTH)
    
    # Extract missing keywords from analysis
    missing_keywords = []
    if analysis_results and "keyword_gaps" in analysis_results:
        gaps = analysis_results["keyword_gaps"]
        for category in ["required_skills", "preferred_skills", "tools_technologies", "soft_skills"]:
            if category in gaps and gaps[category]:
                missing_keywords.extend(gaps[category])
    
    # Build revision prompt
    user_prompt = f"""{create_revision_prompt()}

ORIGINAL CV:
{cv_text_truncated}

JOB DESCRIPTION (for context):
{jd_text_truncated}

MISSING KEYWORDS TO INTEGRATE (naturally, where relevant):
{', '.join(missing_keywords) if missing_keywords else 'None identified'}

ATS MATCH SCORE: {analysis_results.get('match_score', 'N/A')}%

Now provide the complete revised CV text with improvements."""
    
    # Call appropriate API with retry logic
    result = None
    
    try:
        with st.spinner(f"Generating revised CV with {llm_choice}..."):
            if llm_choice == "GPT-4o":
                result = call_openai_revision_api(user_prompt)
            elif llm_choice == "Gemini 1.5 Pro":
                result = call_gemini_revision_api(user_prompt)
        
        # Cache result if successful
        if result:
            st.session_state.revision_cache[cache_key] = result
        
        return result
    
    except Exception as e:
        st.error(f"Error generating revised CV: {str(e)}")
        return None


def call_openai_revision_api(user_prompt, max_retries=MAX_RETRIES):
    """
    Call OpenAI API for CV revision.
    
    Args:
        user_prompt: Complete prompt with CV and instructions
        max_retries: Maximum number of retry attempts
        
    Returns:
        Revised CV text string or None if all retries fail
    """
    try:
        from openai import OpenAI
        
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        for attempt in range(max_retries):
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3,  # Slightly higher for creative revision
                    timeout=API_TIMEOUT
                )
                
                revised_cv = response.choices[0].message.content
                
                # Basic validation - ensure we got substantial content
                if revised_cv and len(revised_cv.strip()) > 100:
                    return revised_cv.strip()
                else:
                    if attempt < max_retries - 1:
                        time.sleep(RETRY_DELAYS[attempt])
                        continue
                    else:
                        st.error("Received insufficient content from API")
                        return None
            
            except Exception as e:
                error_str = str(e)
                # Handle quota errors (don't retry)
                if "quota" in error_str.lower() or "429" in error_str:
                    st.error("OpenAI API quota exceeded. Please check your billing or switch to Gemini.")
                    return None
                
                # For other errors, retry if attempts remain
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAYS[attempt])
                    continue
                else:
                    st.error(f"OpenAI API error: {error_str}")
                    return None
    
    except ImportError:
        st.error("OpenAI library not installed. Run: pip install openai")
        return None
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {str(e)}")
        return None


def call_gemini_revision_api(user_prompt, max_retries=MAX_RETRIES):
    """
    Call Google Gemini API for CV revision.
    
    Args:
        user_prompt: Complete prompt with CV and instructions
        max_retries: Maximum number of retry attempts
        
    Returns:
        Revised CV text string or None if all retries fail
    """
    try:
        import google.generativeai as genai
        
        genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
        
        # Try multiple model names for compatibility
        model_names = [
            "gemini-1.5-pro",
            "gemini-1.5-flash",
            "gemini-pro"
        ]
        
        model = None
        for name in model_names:
            try:
                model = genai.GenerativeModel(
                    model_name=name,
                    generation_config={"temperature": 0.3}
                )
                break
            except Exception:
                continue
        
        if model is None:
            st.error("Could not initialize Gemini model")
            return None
        
        for attempt in range(max_retries):
            try:
                response = model.generate_content(user_prompt)
                revised_cv = response.text
                
                # Basic validation
                if revised_cv and len(revised_cv.strip()) > 100:
                    return revised_cv.strip()
                else:
                    if attempt < max_retries - 1:
                        time.sleep(RETRY_DELAYS[attempt])
                        continue
                    else:
                        st.error("Received insufficient content from API")
                        return None
            
            except Exception as e:
                error_str = str(e)
                # Handle model errors (don't retry)
                if "404" in error_str or "not found" in error_str.lower():
                    st.error(f"Gemini model not found: {error_str}")
                    return None
                
                # For other errors, retry if attempts remain
                if attempt < max_retries - 1:
                    time.sleep(RETRY_DELAYS[attempt])
                    continue
                else:
                    st.error(f"Gemini API error: {error_str}")
                    return None
    
    except ImportError:
        st.error("Google Generative AI library not installed. Run: pip install google-generativeai")
        return None
    except Exception as e:
        st.error(f"Error initializing Gemini client: {str(e)}")
        return None


# ============================================================================
# CV REVISION VALIDATION & DOCX GENERATION
# ============================================================================

def validate_revised_cv(revised_cv_text):
    """
    Validate revised CV text and apply fallback logic if malformed.
    
    Checks for basic structure and content quality. If the LLM output
    is malformed or insufficient, attempts to parse and clean it.
    
    Args:
        revised_cv_text: Revised CV text from LLM
        
    Returns:
        Tuple of (is_valid, cleaned_text, warnings_list)
    """
    warnings = []
    
    # Check if we have content
    if not revised_cv_text or len(revised_cv_text.strip()) < 100:
        return False, None, ["Revised CV is too short or empty"]
    
    cleaned_text = revised_cv_text.strip()
    
    # Check for common section headings (basic structure validation)
    common_headings = [
        "summary", "experience", "education", "skills", "profile",
        "objective", "work", "employment", "qualification", "expertise"
    ]
    
    has_structure = any(heading in cleaned_text.lower() for heading in common_headings)
    
    if not has_structure:
        warnings.append("No standard CV sections detected - output may lack proper structure")
    
    # Check if output contains excessive meta-text or explanations
    meta_phrases = [
        "here is the revised", "i have revised", "the updated cv",
        "as requested", "here's the", "revised version"
    ]
    
    first_100_chars = cleaned_text[:100].lower()
    if any(phrase in first_100_chars for phrase in meta_phrases):
        # Try to remove meta-text from beginning
        lines = cleaned_text.split('\n')
        # Skip first few lines that contain meta-text
        for i, line in enumerate(lines):
            if any(phrase in line.lower() for phrase in meta_phrases):
                continue
            else:
                cleaned_text = '\n'.join(lines[i:])
                warnings.append("Removed meta-text from beginning of output")
                break
    
    # Check for bullet points (indicating preserved structure)
    has_bullets = ('•' in cleaned_text or 
                   '\n-' in cleaned_text or 
                   '\n*' in cleaned_text or
                   cleaned_text.count('- ') > 2)
    
    if not has_bullets:
        warnings.append("No bullet points detected - structure may not be preserved")
    
    # Final validation
    is_valid = len(cleaned_text) >= 100 and (has_structure or has_bullets)
    
    return is_valid, cleaned_text, warnings


def parse_cv_sections(cv_text):
    """
    Parse CV text into structured sections for DOCX generation.
    
    Attempts to identify headings and their content to preserve
    the original structure in the Word document.
    
    Args:
        cv_text: Complete CV text
        
    Returns:
        List of tuples: [(heading, content), ...]
    """
    sections = []
    lines = cv_text.split('\n')
    
    current_heading = None
    current_content = []
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            # Empty line - might be section separator
            if current_content and current_content[-1] != '':
                current_content.append('')
            continue
        
        # Heuristic: A heading is typically:
        # - Short (< 50 chars)
        # - ALL CAPS or Title Case
        # - No bullet points
        # - Not too many words (< 6 words)
        is_heading = (
            len(stripped) < 50 and
            not stripped.startswith(('-', '•', '*', '–')) and
            len(stripped.split()) <= 6 and
            (stripped.isupper() or stripped.istitle())
        )
        
        if is_heading and current_heading is not None:
            # Save previous section
            sections.append((current_heading, '\n'.join(current_content)))
            current_heading = stripped
            current_content = []
        elif is_heading:
            # First heading
            current_heading = stripped
            current_content = []
        else:
            # Regular content
            if current_heading is None:
                # Content before first heading
                current_heading = "Profile"
            current_content.append(line)  # Preserve original indentation
    
    # Save last section
    if current_heading:
        sections.append((current_heading, '\n'.join(current_content)))
    
    return sections


def set_cell_margins(cell, **kwargs):
    """
    Set cell margins (padding) in twips (1/1440 inch) using XML manipulation.
    
    Args:
        cell: Table cell object
        **kwargs: top, bottom, start (left), end (right) in twips
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        
        for margin_name, margin_value in kwargs.items():
            if margin_value is not None:
                node = OxmlElement(f'w:{margin_name}')
                node.set(qn('w:w'), str(margin_value))
                node.set(qn('w:type'), 'dxa')  # twips
                tcMar.append(node)
        
        tcPr.append(tcMar)
    except Exception:
        pass  # Silently fail if XML manipulation fails


def remove_table_borders(table):
    """
    Remove all table borders using XML manipulation to create "invisible" table.
    
    This is the ONLY way to truly remove borders in python-docx.
    Setting table.borders = False is insufficient.
    
    Args:
        table: Document table object
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Iterate through all cells to set border properties
        for row in table.rows:
            for cell in row.cells:
                # Get or create table cell properties (tcPr)
                tcPr = cell._element.get_or_add_tcPr()
                
                # Create table cell borders element (tcBorders)
                tcBorders = OxmlElement('w:tcBorders')
                
                # Set all six border types to 'nil' (no border)
                for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    border_el = OxmlElement(f'w:{edge}')
                    border_el.set(qn('w:val'), 'nil')
                    border_el.set(qn('w:sz'), '0')
                    border_el.set(qn('w:space'), '0')
                    border_el.set(qn('w:color'), 'auto')
                    tcBorders.append(border_el)
                
                # Append borders to cell properties
                tcPr.append(tcBorders)
    except Exception as e:
        st.warning(f"Could not remove table borders (non-critical): {str(e)}")


def parse_cv_structure_for_docx(cv_text):
    """
    Parse CV text into structured components for DOCX generation.
    
    Identifies:
    - Contact information (name, email, phone, location)
    - Skills (grouped by category)
    - Experience entries
    - Projects
    - Education
    
    Args:
        cv_text: Complete CV text
        
    Returns:
        dict with structured content
    """
    structure = {
        'contact': [],
        'skills': {},
        'experience': [],
        'projects': [],
        'education': [],
        'other_sections': []
    }
    
    lines = cv_text.split('\n')
    current_section = None
    current_content = []
    
    # Common section headers
    skills_keywords = ['skill', 'technical', 'technologies', 'expertise', 'competencies']
    experience_keywords = ['experience', 'employment', 'work history', 'professional']
    project_keywords = ['project', 'portfolio']
    education_keywords = ['education', 'qualification', 'academic']
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Check if line is a section header (heuristic)
        is_header = (len(stripped) < 50 and 
                    not stripped.startswith(('-', '•', '*', '–', '◦')) and
                    (stripped.isupper() or stripped.istitle()))
        
        if is_header:
            # Save previous section
            if current_section and current_content:
                if current_section == 'skills':
                    structure['skills'] = parse_skills_section(current_content)
                elif current_section == 'experience':
                    structure['experience'] = parse_experience_section(current_content)
                elif current_section == 'projects':
                    structure['projects'] = parse_projects_section(current_content)
                elif current_section == 'education':
                    structure['education'] = parse_education_section(current_content)
                else:
                    structure['other_sections'].append((current_section, '\n'.join(current_content)))
            
            # Identify new section
            lower_header = stripped.lower()
            if any(kw in lower_header for kw in skills_keywords):
                current_section = 'skills'
            elif any(kw in lower_header for kw in experience_keywords):
                current_section = 'experience'
            elif any(kw in lower_header for kw in project_keywords):
                current_section = 'projects'
            elif any(kw in lower_header for kw in education_keywords):
                current_section = 'education'
            else:
                current_section = stripped
            
            current_content = []
        else:
            current_content.append(stripped)
    
    # Save last section
    if current_section and current_content:
        if current_section == 'skills':
            structure['skills'] = parse_skills_section(current_content)
        elif current_section == 'experience':
            structure['experience'] = parse_experience_section(current_content)
        elif current_section == 'projects':
            structure['projects'] = parse_projects_section(current_content)
        elif current_section == 'education':
            structure['education'] = parse_education_section(current_content)
        else:
            structure['other_sections'].append((current_section, '\n'.join(current_content)))
    
    return structure


def parse_skills_section(lines):
    """Parse skills into grouped categories."""
    skills_dict = {}
    current_group = "General"
    
    for line in lines:
        # Check if line is a group header (e.g., "Frontend:", "Backend")
        if ':' in line and len(line.split(':')[0].split()) <= 3:
            current_group = line.split(':')[0].strip()
            skills_text = line.split(':', 1)[1].strip()
            if skills_text:
                skills_dict[current_group] = [s.strip() for s in skills_text.split(',')]
        elif not line.startswith(('-', '•', '*')):
            # Might be a group header without colon
            if len(line.split()) <= 3 and line[0].isupper():
                current_group = line.strip()
                skills_dict[current_group] = []
            else:
                # Skills list
                if current_group not in skills_dict:
                    skills_dict[current_group] = []
                skills_dict[current_group].extend([s.strip() for s in line.split(',')])
        else:
            # Bullet point format
            skill = line.lstrip('-•*◦ ').strip()
            if current_group not in skills_dict:
                skills_dict[current_group] = []
            skills_dict[current_group].append(skill)
    
    return skills_dict


def parse_experience_section(lines):
    """Parse experience entries with title, company, dates, bullets."""
    entries = []
    current_entry = None
    
    for line in lines:
        if line.startswith(('-', '•', '*', '–', '◦')):
            # Bullet point
            if current_entry:
                current_entry['bullets'].append(line.lstrip('-•*–◦ ').strip())
        else:
            # Potential job title/company or date
            if current_entry and 'dates' not in current_entry:
                # Check if this line contains dates
                if re.search(r'\d{4}|present|current', line, re.IGNORECASE):
                    current_entry['dates'] = line
                else:
                    # Part of title or company
                    current_entry['title'] += ' ' + line
            else:
                # Start new entry
                if current_entry:
                    entries.append(current_entry)
                current_entry = {
                    'title': line,
                    'company': '',
                    'dates': '',
                    'bullets': []
                }
    
    if current_entry:
        entries.append(current_entry)
    
    return entries


def parse_projects_section(lines):
    """Parse projects with title, tech stack, and bullets."""
    projects = []
    current_project = None
    
    for line in lines:
        if line.startswith(('-', '•', '*', '–', '◦')):
            # Bullet point
            if current_project:
                current_project['bullets'].append(line.lstrip('-•*–◦ ').strip())
        else:
            # Check if tech stack line
            if 'tech' in line.lower() and ':' in line:
                if current_project:
                    current_project['tech_stack'] = line.split(':', 1)[1].strip()
            else:
                # New project title
                if current_project:
                    projects.append(current_project)
                current_project = {
                    'title': line,
                    'tech_stack': '',
                    'bullets': []
                }
    
    if current_project:
        projects.append(current_project)
    
    return projects


def parse_education_section(lines):
    """Parse education entries."""
    return [line for line in lines if line.strip()]


def generate_professional_docx(revised_cv_text, layout_info=None):
    """
    Generate a professional Word document with invisible 2-column table layout.
    
    Creates an ATS-friendly .docx file with:
    - Invisible 2-column table (30% sidebar, 70% main)
    - Professional fonts (Arial, 11pt body, 14pt headings)
    - Narrow margins (0.5 inches)
    - Colors: Headers #000000, Body #333333
    - Proper spacing and hierarchy
    
    Args:
        revised_cv_text: Complete revised CV text
        layout_info: Optional dict with layout metadata (from extraction)
        
    Returns:
        BytesIO object containing the .docx file, or None if generation fails
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Create new document
        doc = Document()
        
        # Set document margins (0.5 inches - Narrow preset)
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Parse CV into structured components
        cv_structure = parse_cv_structure_for_docx(revised_cv_text)
        
        # Determine if we should use 2-column layout
        use_two_column = (layout_info and layout_info.get('layout') == '2-col') or bool(cv_structure['skills'])
        
        if use_two_column and cv_structure['skills']:
            # Create invisible 2-column table
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Set column widths (30% / 70%)
            table.columns[0].width = Cm(6.0)   # ~2.36 inches (sidebar)
            table.columns[1].width = Cm(14.0)  # ~5.51 inches (main)
            
            # Remove all table borders (invisible table)
            remove_table_borders(table)
            
            # Get cell references
            sidebar_cell = table.cell(0, 0)
            main_cell = table.cell(0, 1)
            
            # Set cell padding
            set_cell_margins(sidebar_cell, top=100, bottom=100, start=100, end=200)
            set_cell_margins(main_cell, top=100, bottom=100, start=200, end=100)
            
            # Populate sidebar (Skills + Contact)
            if cv_structure['skills']:
                p = sidebar_cell.add_paragraph()
                run = p.add_run("SKILLS")
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                
                for skill_group, skills_list in cv_structure['skills'].items():
                    # Group header
                    p = sidebar_cell.add_paragraph()
                    run = p.add_run(skill_group)
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(2)
                    
                    # Skills (comma-separated or bullets)
                    if skills_list:
                        skills_text = ', '.join(skills_list) if len(skills_list) > 1 else skills_list[0]
                        p = sidebar_cell.add_paragraph(skills_text)
                        p.style.font.size = Pt(9)
                        p.style.font.color.rgb = RGBColor(51, 51, 51)
                        p.paragraph_format.space_after = Pt(3)
            
            # Populate main cell (Experience + Projects + Education)
            if cv_structure['experience']:
                p = main_cell.add_paragraph()
                run = p.add_run("PROFESSIONAL EXPERIENCE")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.paragraph_format.space_after = Pt(6)
                
                for exp_entry in cv_structure['experience']:
                    # Job title
                    p = main_cell.add_paragraph()
                    run = p.add_run(exp_entry['title'])
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.space_before = Pt(6)
                    
                    # Dates
                    if exp_entry.get('dates'):
                        p = main_cell.add_paragraph(exp_entry['dates'])
                        p.style.font.size = Pt(10)
                        p.style.font.italic = True
                        p.style.font.color.rgb = RGBColor(102, 102, 102)
                    
                    # Bullets
                    for bullet in exp_entry['bullets']:
                        p = main_cell.add_paragraph(bullet, style='List Bullet')
                        p.style.font.size = Pt(11)
                        p.style.font.color.rgb = RGBColor(51, 51, 51)
                        p.paragraph_format.space_after = Pt(3)
            
            if cv_structure['projects']:
                p = main_cell.add_paragraph()
                run = p.add_run("PROJECTS")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                
                for project in cv_structure['projects']:
                    # Project title
                    p = main_cell.add_paragraph()
                    run = p.add_run(project['title'])
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.space_before = Pt(6)
                    
                    # Tech stack
                    if project.get('tech_stack'):
                        p = main_cell.add_paragraph()
                        run = p.add_run(f"Tech Stack: {project['tech_stack']}")
                        run.font.size = Pt(10)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(102, 102, 102)
                    
                    # Bullets
                    for bullet in project['bullets']:
                        p = main_cell.add_paragraph(bullet, style='List Bullet')
                        p.style.font.size = Pt(11)
                        p.style.font.color.rgb = RGBColor(51, 51, 51)
                        p.paragraph_format.space_after = Pt(3)
            
            if cv_structure['education']:
                p = main_cell.add_paragraph()
                run = p.add_run("EDUCATION")
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                
                for edu_line in cv_structure['education']:
                    p = main_cell.add_paragraph(edu_line)
                    p.style.font.size = Pt(11)
                    p.style.font.color.rgb = RGBColor(51, 51, 51)
        
        else:
            # Fallback: Simple 1-column layout
            cv_sections = parse_cv_sections(revised_cv_text)
            
            if not cv_sections:
                lines = revised_cv_text.split('\n')
                for line in lines:
                    if line.strip():
                        p = doc.add_paragraph(line)
                        run = p.runs[0] if p.runs else p.add_run()
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
            else:
                for i, (heading, content) in enumerate(cv_sections):
                    # Add section heading
                    heading_para = doc.add_paragraph(heading)
                    run = heading_para.runs[0] if heading_para.runs else heading_para.add_run()
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    heading_para.paragraph_format.space_after = Pt(6)
                    
                    # Process content
                    content_lines = content.split('\n')
                    for line in content_lines:
                        stripped = line.strip()
                        if not stripped:
                            continue
                        
                        is_bullet = stripped.startswith(('-', '•', '*', '–', '◦'))
                        
                        if is_bullet:
                            text = stripped.lstrip('-•*–◦ ').strip()
                            p = doc.add_paragraph(text, style='List Bullet')
                            run = p.runs[0] if p.runs else p.add_run()
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            p.paragraph_format.space_after = Pt(3)
                        else:
                            p = doc.add_paragraph(stripped)
                            run = p.runs[0] if p.runs else p.add_run()
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            p.paragraph_format.space_after = Pt(6)
                    
                    if i < len(cv_sections) - 1:
                        doc.add_paragraph()
        
        # Save to BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    except ImportError:
        st.error("python-docx library not installed. Run: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Error generating Word document: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None


# ============================================================================
# UI RENDERING
# ============================================================================

def render_score_breakdown(breakdown):
    """
    Render score breakdown visualization.
    
    Displays each scoring component with its value and percentage of maximum.
    Uses color-coded metrics: green (high), yellow (medium), red (low).
    
    Args:
        breakdown: Dictionary with technical_skills, experience_alignment, keyword_coverage
    """
    st.subheader("📊 Score Breakdown")
    
    # Create three columns for the metrics
    col1, col2, col3 = st.columns(3)
    
    with col1:
        tech_score = breakdown.get("technical_skills", 0)
        tech_pct = (tech_score / 50) * 100
        delta_color = "normal" if tech_score >= 35 else "inverse"
        st.metric(
            label="Technical Skills",
            value=f"{tech_score}/50",
            delta=f"{tech_pct:.0f}%"
        )
    
    with col2:
        exp_score = breakdown.get("experience_alignment", 0)
        exp_pct = (exp_score / 30) * 100
        st.metric(
            label="Experience",
            value=f"{exp_score}/30",
            delta=f"{exp_pct:.0f}%"
        )
    
    with col3:
        kw_score = breakdown.get("keyword_coverage", 0)
        kw_pct = (kw_score / 20) * 100
        st.metric(
            label="Keywords",
            value=f"{kw_score}/20",
            delta=f"{kw_pct:.0f}%"
        )


def render_keyword_gaps(keyword_gaps):
    """
    Render keyword gap analysis section.
    
    Displays keywords from JD that are missing or weak in the CV,
    grouped by category (required skills, preferred skills, tools, soft skills).
    
    Args:
        keyword_gaps: Dictionary with categorized keyword lists
    """
    with st.expander("🔍 Keyword Gap Analysis", expanded=True):
        # Check if any gaps exist
        has_gaps = False
        for category in keyword_gaps.values():
            if category:
                has_gaps = True
                break
        
        if not has_gaps:
            st.success("✅ No significant keyword gaps detected!")
            return
        
        # Required Skills
        if keyword_gaps.get("required_skills"):
            st.markdown("**⚠️ Required Skills (Missing/Weak):**")
            for skill in keyword_gaps["required_skills"]:
                st.markdown(f"- {skill}")
            st.markdown("")
        
        # Preferred Skills
        if keyword_gaps.get("preferred_skills"):
            st.markdown("**💡 Preferred Skills (Missing/Weak):**")
            for skill in keyword_gaps["preferred_skills"]:
                st.markdown(f"- {skill}")
            st.markdown("")
        
        # Tools & Technologies
        if keyword_gaps.get("tools_technologies"):
            st.markdown("**🛠️ Tools & Technologies (Missing/Weak):**")
            for tool in keyword_gaps["tools_technologies"]:
                st.markdown(f"- {tool}")
            st.markdown("")
        
        # Soft Skills
        if keyword_gaps.get("soft_skills"):
            st.markdown("**🤝 Soft Skills (Missing/Weak):**")
            for skill in keyword_gaps["soft_skills"]:
                st.markdown(f"- {skill}")


def render_bullet_optimizations(optimizations):
    """
    Render bullet point optimization section.
    
    Shows original bullets alongside optimized versions using X-Y-Z formula,
    with explanations of why the changes improve ATS matching.
    
    Args:
        optimizations: List of optimization dictionaries
    """
    with st.expander("✍️ Bullet Point Optimizations (X-Y-Z Formula)", expanded=False):
        if not optimizations:
            st.info("No bullet point optimizations suggested.")
            return
        
        for i, opt in enumerate(optimizations, 1):
            st.markdown(f"**Bullet {i}:**")
            
            # Original
            st.markdown("*Original:*")
            st.markdown(f"> {opt.get('original', 'N/A')}")
            
            # Optimized
            st.markdown("*Optimized (X-Y-Z Formula):*")
            st.success(opt.get('optimized', 'N/A'))
            
            # Reason
            st.markdown(f"*Why:* {opt.get('reason', 'N/A')}")
            st.markdown("---")


def render_ats_red_flags(red_flags):
    """
    Render ATS red flags section.
    
    Displays detected ATS parsing issues with explanations and actionable fixes.
    Uses error styling for critical issues.
    
    Args:
        red_flags: List of red flag dictionaries
    """
    with st.expander("🚨 ATS Red Flags & Fixes", expanded=True):
        if not red_flags:
            st.success("✅ No critical ATS red flags detected!")
            return
        
        for i, flag in enumerate(red_flags, 1):
            st.markdown(f"**Issue {i}: {flag.get('issue', 'Unknown issue')}**")
            
            # Explanation
            st.markdown(f"*Why ATS struggles:* {flag.get('explanation', 'N/A')}")
            
            # Fix
            st.markdown("*✅ Fix:*")
            st.info(flag.get('fix', 'N/A'))
            
            st.markdown("---")


def render_cv_revision_section(cv_text, jd_text, analysis_results, llm_choice):
    """
    Render enhanced CV revision section with validation feedback and metrics.
    
    Allows users to generate a professionally revised CV based on
    analysis feedback, with structure validation, metric checking,
    and professional DOCX with invisible table layout.
    
    Args:
        cv_text: Original CV text
        jd_text: Job description text
        analysis_results: Analysis results dictionary
        llm_choice: Selected LLM model
    """
    st.markdown("---")
    
    with st.expander("✏️ Generate Revised CV", expanded=False):
        st.markdown("""
        **Generate a professionally revised CV with:**
        - ✨ X-Y-Z formula with quantifiable metrics
        - 🔧 Automatic date error correction
        - 🎯 Missing keywords naturally integrated
        - 📐 Professional 2-column layout (if applicable)
        - ✅ Structure preservation validation
        
        ⚠️ **Note**: Graphics and images cannot be preserved in ATS-friendly DOCX format.
        """)
        
        # Performance warning for large CVs
        if len(cv_text) > 15000:
            st.warning("⚠️ Your CV is quite long. Revision may take longer than usual.")
        
        # Generate button with updated text
        generate_button = st.button("✨ Build My Revised CV (.docx)", type="primary", key="generate_revision")
        
        if generate_button:
            progress_placeholder = st.empty()
            validation_placeholder = st.empty()
            
            # Step 1: Analyze layout
            progress_placeholder.info("🔍 Analyzing layout structure (detecting columns)...")
            layout_info = None
            try:
                # Get layout info from session if available
                if "last_pdf_bytes" in st.session_state:
                    layout_info = extract_text_with_layout(st.session_state.last_pdf_bytes)
                    if layout_info.get('layout') == '2-col':
                        progress_placeholder.success(f"✅ Layout detected: 2-column (sidebar at x={layout_info.get('split_x', 'N/A'):.0f}px)")
                    else:
                        progress_placeholder.info("ℹ️ Layout detected: 1-column")
                    time.sleep(0.5)
            except Exception:
                pass
            
            # Step 2: Fix dates
            progress_placeholder.info("📅 Fixing logical date errors...")
            corrected_cv, date_corrections = fix_future_dates(cv_text)
            if date_corrections:
                validation_placeholder.success(f"✅ Fixed {len(date_corrections)} future date error(s): {', '.join(date_corrections[:2])}")
            time.sleep(0.5)
            
            # Step 3: LLM Revision
            progress_placeholder.info("✨ Applying X-Y-Z formula improvements...")
            revised_cv = call_revision_llm(corrected_cv, jd_text, analysis_results, llm_choice)
            
            if revised_cv:
                # Step 4: Validate structure
                progress_placeholder.info("📊 Validating structure preservation...")
                is_valid, cleaned_cv, warnings = validate_revised_cv(revised_cv)
                
                # Step 5: Validate metrics
                progress_placeholder.info("🔢 Checking quantifiable metrics...")
                metric_validation = validate_quantifiable_bullets(cleaned_cv if cleaned_cv else revised_cv)
                
                bullets_with_metrics = sum(1 for _, has_metric, _ in metric_validation if has_metric)
                bullets_without_metrics = sum(1 for _, has_metric, _ in metric_validation if not has_metric)
                
                progress_placeholder.empty()
                
                # Show validation results
                st.markdown("#### 📊 Validation Results")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if date_corrections:
                        st.metric("Date Fixes", len(date_corrections), delta="corrected", delta_color="normal")
                    else:
                        st.metric("Date Fixes", "0", delta="none needed")
                
                with col2:
                    if bullets_with_metrics > 0:
                        st.metric("Metrics Applied", bullets_with_metrics, delta="quantifiable")
                    else:
                        st.metric("Metrics Applied", "0")
                
                with col3:
                    if bullets_without_metrics > 0:
                        st.metric("Needs Metrics", bullets_without_metrics, delta="review needed", delta_color="inverse")
                    else:
                        st.metric("Needs Metrics", "0", delta="all good")
                
                # Show warnings if any
                if warnings:
                    with st.expander("⚠️ Validation Warnings", expanded=True):
                        for warning in warnings:
                            st.warning(warning)
                
                if is_valid and cleaned_cv:
                    # Store in session state
                    st.session_state.revised_cv = cleaned_cv
                    st.session_state.layout_info = layout_info
                    st.session_state.metric_validation = metric_validation
                    st.success("✅ Revised CV generated successfully!")
                else:
                    st.error("❌ Failed to generate valid revised CV. Please try again.")
                    return
            else:
                progress_placeholder.empty()
                st.error("❌ Failed to generate revised CV. Please check your API key and try again.")
                return
        
        # Display preview and download if revised CV exists in session
        if "revised_cv" in st.session_state:
            st.markdown("---")
            st.markdown("### 📄 Revised CV Preview")
            st.info("*Preview shows content structure. Bullets flagged with [NEEDS METRIC] require quantifiable data.*")
            
            # Display preview with metric validation
            revised_cv = st.session_state.revised_cv
            metric_validation = st.session_state.get('metric_validation', [])
            
            # Create a mapping of bullets to their validation status
            bullet_status = {bullet_text: (has_metric, warning) for bullet_text, has_metric, warning in metric_validation}
            
            # Parse into sections for better preview
            sections = parse_cv_sections(revised_cv)
            
            if sections:
                # Display as structured sections with metric flags
                for heading, content in sections:
                    st.markdown(f"### {heading}")
                    
                    # Format content with bullets and metric flags
                    content_lines = content.split('\n')
                    for line in content_lines:
                        stripped = line.strip()
                        if stripped:
                            if stripped.startswith(('-', '•', '*', '–', '◦')):
                                # Display as bullet with metric validation
                                text = stripped.lstrip('-•*–◦ ').strip()
                                
                                # Check if this bullet needs metrics
                                needs_metric = False
                                for bullet_text, (has_metric, warning) in bullet_status.items():
                                    if bullet_text in text or text in bullet_text:
                                        if not has_metric:
                                            needs_metric = True
                                        break
                                
                                if needs_metric:
                                    st.markdown(f"- ⚠️ **[NEEDS METRIC]** {text}")
                                else:
                                    st.markdown(f"- {text}")
                            else:
                                # Display as paragraph
                                st.markdown(stripped)
                    
                    st.markdown("")  # Add spacing between sections
            else:
                # Fallback: display as plain text
                st.text_area("Revised CV", revised_cv, height=400, key="preview_text")
            
            # Expandable validation details
            if metric_validation:
                with st.expander("🔍 Detailed Metric Validation", expanded=False):
                    st.markdown("**Bullet Points Analysis:**")
                    for bullet_text, has_metric, warning in metric_validation:
                        if has_metric:
                            st.success(f"✅ {bullet_text[:80]}...")
                        else:
                            st.warning(warning)
            
            st.markdown("---")
            
            # Generate DOCX for download
            st.markdown("### 💾 Download Revised CV")
            st.info("📄 Generating professional DOCX with invisible 2-column table layout...")
            
            layout_info = st.session_state.get('layout_info', None)
            docx_buffer = generate_professional_docx(revised_cv, layout_info)
            
            if docx_buffer:
                st.download_button(
                    label="📥 Download as Word Document (.docx)",
                    data=docx_buffer,
                    file_name="revised_cv_professional.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_docx"
                )
                st.success("✅ Your professional revised CV is ready for download!")
                
                if layout_info and layout_info.get('layout') == '2-col':
                    st.info("📐 Generated with 2-column layout (30% sidebar / 70% main content)")
            else:
                st.error("❌ Failed to generate Word document. Please try again.")


def render_analysis_results(data):
    """
    Main results rendering function.
    
    Renders all analysis sections with defensive checks to handle partial
    or malformed data gracefully. Each section checks for data existence
    before attempting to render.
    
    Args:
        data: Complete analysis results dictionary
    """
    if not data:
        st.error("No analysis results to display.")
        return
    
    # Match Score (main metric)
    match_score = data.get("match_score", 0)
    st.metric(
        label="ATS Match Score",
        value=f"{match_score}%",
        help="Overall match score based on technical skills (50%), experience (30%), and keywords (20%)"
    )
    
    # Progress bar visualization
    if match_score >= 70:
        st.progress(match_score / 100)
        st.success(f"🎉 Strong match! Your CV aligns well with the job requirements.")
    elif match_score >= 50:
        st.progress(match_score / 100)
        st.warning(f"⚠️ Moderate match. Consider addressing the gaps below.")
    else:
        st.progress(match_score / 100)
        st.error(f"❌ Weak match. Significant improvements needed.")
    
    # Score explanation
    if data.get("score_explanation"):
        st.markdown("**Score Explanation:**")
        st.info(data["score_explanation"])
    
    st.markdown("---")
    
    # Score breakdown visualization
    if data.get("score_breakdown"):
        render_score_breakdown(data["score_breakdown"])
        st.markdown("---")
    
    # Keyword gaps
    if data.get("keyword_gaps"):
        render_keyword_gaps(data["keyword_gaps"])
        st.markdown("---")
    
    # Bullet optimizations
    if data.get("bullet_optimizations"):
        render_bullet_optimizations(data["bullet_optimizations"])
        st.markdown("---")
    
    # ATS red flags
    if data.get("ats_red_flags"):
        render_ats_red_flags(data["ats_red_flags"])


# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    """
    Main Streamlit application entry point.
    
    Sets up the UI layout, handles user interactions, and orchestrates
    the analysis workflow from PDF upload to results display.
    """
    # Page configuration
    st.set_page_config(
        page_title="Smart CV Optimizer",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Title and description
    st.title("📄 Smart CV Optimizer")
    st.markdown("""
    Optimize your CV to pass Applicant Tracking Systems (ATS) and improve alignment with Job Descriptions.
    Upload your CV, paste the job description, and get actionable insights.
    """)
    
    # Sidebar configuration
    st.sidebar.header("⚙️ Configuration")
    
    # LLM selection
    llm_choice = st.sidebar.selectbox(
        "Select LLM",
        ["GPT-4o", "Gemini 1.5 Pro"],
        help="Choose the language model for analysis"
    )
    
    # Check API keys
    st.sidebar.markdown("---")
    st.sidebar.markdown("**API Status:**")
    
    if llm_choice == "GPT-4o":
        if os.getenv("OPENAI_API_KEY"):
            st.sidebar.success("✅ OpenAI API key detected")
        else:
            st.sidebar.error("❌ OpenAI API key not found")
            st.sidebar.info("Set OPENAI_API_KEY in your .env file")
    
    elif llm_choice == "Gemini 1.5 Pro":
        if os.getenv("GOOGLE_API_KEY"):
            st.sidebar.success("✅ Google API key detected")
        else:
            st.sidebar.error("❌ Google API key not found")
            st.sidebar.info("Set GOOGLE_API_KEY in your .env file")
    
    # Main layout: two columns
    left_col, right_col = st.columns([1, 1])
    
    # ========================================================================
    # LEFT COLUMN: Input Section
    # ========================================================================
    with left_col:
        st.header("📥 Input")
        
        # PDF Upload
        uploaded_file = st.file_uploader(
            "Upload your CV (PDF)",
            type=["pdf"],
            help="Upload your CV in PDF format"
        )
        
        # Job Description Input
        jd_text = st.text_area(
            "Paste Job Description",
            height=300,
            placeholder="Paste the complete job description here...",
            help="Paste the full job description text"
        )
        
        # Show character counts
        if jd_text:
            jd_char_count = len(jd_text)
            if jd_char_count > MAX_JD_LENGTH:
                st.warning(f"⚠️ Job Description: {jd_char_count} characters (will be truncated to {MAX_JD_LENGTH})")
            else:
                st.info(f"📊 Job Description: {jd_char_count} characters")
        
        # Analyze button
        analyze_button = st.button("🚀 Analyze CV", type="primary", use_container_width=True)
    
    # ========================================================================
    # RIGHT COLUMN: Results Section
    # ========================================================================
    with right_col:
        st.header("📊 Analysis Results")
        
        # Process analysis when button is clicked
        if analyze_button:
            # Validation
            if not uploaded_file:
                st.error("❌ Please upload a CV in PDF format.")
                return
            
            if not jd_text or len(jd_text.strip()) < 50:
                st.error("❌ Please provide a valid job description (at least 50 characters).")
                return
            
            # Check API key
            if llm_choice == "GPT-4o" and not os.getenv("OPENAI_API_KEY"):
                st.error("❌ OpenAI API key not found. Please set OPENAI_API_KEY in your .env file.")
                return
            
            if llm_choice == "Gemini 1.5 Pro" and not os.getenv("GOOGLE_API_KEY"):
                st.error("❌ Google API key not found. Please set GOOGLE_API_KEY in your .env file.")
                return
            
            # Extract text from PDF
            with st.spinner("Extracting text from PDF..."):
                pdf_bytes = uploaded_file.read()
                cv_text = extract_text_from_pdf(pdf_bytes)
                
                # Store PDF bytes for layout extraction later
                st.session_state.last_pdf_bytes = pdf_bytes
            
            if not cv_text:
                st.error("❌ Failed to extract text from PDF. Please ensure the PDF is readable.")
                return
            
            # Show CV character count
            cv_char_count = len(cv_text)
            if cv_char_count > MAX_CV_LENGTH:
                st.warning(f"⚠️ CV: {cv_char_count} characters (will be truncated to {MAX_CV_LENGTH})")
            else:
                st.info(f"📊 CV: {cv_char_count} characters")
            
            # Call LLM API
            results = call_llm_api(cv_text, jd_text.strip(), llm_choice)
            
            if results:
                # Store results in session state
                st.session_state.last_results = results
                st.session_state.last_cv_text = cv_text
                st.session_state.last_jd_text = jd_text.strip()
                
                # Render results
                render_analysis_results(results)
                
                # Render CV revision section
                render_cv_revision_section(cv_text, jd_text.strip(), results, llm_choice)
            else:
                st.error("❌ Analysis failed. Please check your API key and try again.")
        
        # Display cached results if available (for reruns without re-analysis)
        elif "last_results" in st.session_state:
            st.info("📌 Showing cached results. Click 'Analyze CV' to run a new analysis.")
            render_analysis_results(st.session_state.last_results)
            
            # Render CV revision section with cached data
            if "last_cv_text" in st.session_state and "last_jd_text" in st.session_state:
                render_cv_revision_section(
                    st.session_state.last_cv_text,
                    st.session_state.last_jd_text,
                    st.session_state.last_results,
                    llm_choice
                )
        else:
            st.info("👆 Upload your CV and paste a job description to get started.")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666;'>
        <p>Built with Streamlit | Powered by GPT-4o & Gemini 1.5 Pro</p>
        <p>💡 Tip: Use the X-Y-Z formula to make your achievements more impactful!</p>
    </div>
    """, unsafe_allow_html=True)


# ============================================================================
# APPLICATION ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    main()

